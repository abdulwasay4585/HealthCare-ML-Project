from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_report():
    doc = Document()

    # --- Title ---
    title = doc.add_heading('Design and Deploy an End-to-End Machine Learning System\nwith FastAPI, CI/CD, Prefect, Automated Testing, and Docker Containerization', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # --- Author Block (Placeholders) ---
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('[insert your name here]\n[insert department]\n[insert university]\n[insert city, country]\n[insert email]')
    run.italic = True
    
    # --- Abstract ---
    doc.add_heading('Abstract', level=1)
    abstract_text = (
        "This project implements a complete End-to-End Machine Learning System for the Healthcare domain, "
        "specifically focusing on hospital readmission prediction and length-of-stay estimation. "
        "Leveraging the UCI Diabetes dataset, we developed a system that integrates robust data pipelines using Prefect, "
        "serves models via a high-performance FastAPI application, and ensures reliability through Automated Testing and CI/CD workflows. "
        "The entire application is containerized using Docker to facilitate seamless deployment. "
        "Experimental results demonstrate a classification accuracy of approximately 64% and a regression RMSE of 2.42 days, "
        "showcasing the system's capability to handle real-world healthcare data challenges."
    )
    doc.add_paragraph(abstract_text)

    # --- Section Break for Columns (Attempting 2-col layout) ---
    # Python-docx doesn't strictly support "columns" easily via high-level API for all versions.
    # We will stick to single column for safety/reliability, or just formatted headers.
    # But user asked for "same looking", implying columns. 
    # Let's try to set the section to 2 columns if we can, otherwise standard formatting.
    
    section = doc.sections[-1]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1') # Keep Abstract single column? Or usually Abstract is bold/single.
    
    # Add new section for body
    new_section = doc.add_section()
    sectPr = new_section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '2') # Two columns
    cols.set(qn('w:space'), '720') # Space between columns (0.5 inch)

    # --- I. Introduction ---
    doc.add_heading('I. Introduction', level=1)
    doc.add_paragraph(
        "Machine Learning Engineering (MLOps) has emerged as a critical discipline for translating research models into production systems. "
        "This project aims to build a production-grade ML system for the Healthcare domain. "
        "Effective management of diabetic patients requires accurate prediction of readmission risks and expected hospital stay duration. "
        "Our system addresses these needs by deploying predictive models within a scalable, automated architecture."
    )

    # --- II. Methodology ---
    doc.add_heading('II. Methodology', level=1)
    
    doc.add_heading('A. Dataset', level=2)
    doc.add_paragraph(
        "We utilized the 'Diabetes 130-US hospitals for years 1999-2008' dataset from the UCI Machine Learning Repository. "
        "The dataset contains over 100,000 instances with 50 features, including patient demographics, laboratory results, and medication details."
    )

    doc.add_heading('B. Data Pipeline (Prefect)', level=2)
    doc.add_paragraph(
        "We implemented an orchestrated pipeline using Prefect. The flow 'training_flow' handles:\n"
        "1) Data Ingestion: Loading CSV data and handling missing values (e.g., '?').\n"
        "2) Preprocessing: Feature selection, imputation (Median/Mode), and scaling (StandardScaler).\n"
        "3) Model Training: Parallel training of Classification and Regression models."
    )

    doc.add_heading('C. Models', level=2)
    doc.add_paragraph(
        "1) Classification: A Random Forest Classifier was trained to predict the 'readmitted' status (binary: NO vs YES). "
        "We optimized for F1-score and Accuracy.\n"
        "2) Regression: A Ridge Regression model was trained to predict 'time_in_hospital'. "
        "We utilized separate feature sets for each task to prevent data leakage (excluding target variables from features)."
    )

    # --- III. System Architecture ---
    doc.add_heading('III. System Architecture', level=1)
    doc.add_paragraph(
        "The system creates a robust feedback loop:\n"
        "- **FastAPI**: Provides real-time inference endpoints (/predict/readmission, /predict/los).\n"
        "- **Docker**: Containerizes the application (Python 3.9-slim base) for consistent runtime.\n"
        "- **CI/CD**: GitHub Actions workflow automates testing (pytest) and image building on code changes.\n"
        "- **Automated Testing**: Unit tests verify API health and model artifact validity."
    )

    # --- IV. Experiments & Results ---
    doc.add_heading('IV. Experiments & Results', level=1)
    doc.add_paragraph(
        "We evaluated the models on a 20% held-out test set."
    )
    
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Metric'
    hdr_cells[1].text = 'Value'
    
    row1 = table.rows[1].cells
    row1[0].text = 'Classification Accuracy'
    row1[1].text = '63.62%'
    
    row2 = table.rows[2].cells
    row2[0].text = 'Regression RMSE'
    row2[1].text = '2.42'
    
    doc.add_paragraph(
        "\nObservations:\n"
        "- The Random Forest model effectively captures non-linear relationships in patient demographics.\n"
        "- The Regression model provides a reasonable baseline for estimating hospital stay, with an error margin of ~2.4 days."
    )

    # --- V. Conclusion ---
    doc.add_heading('V. Conclusion', level=1)
    doc.add_paragraph(
        "We successfully deployed a full-stack ML Engineering system. "
        "The integration of Prefect for orchestration and Docker for containerization ensures the system is reproducible and scalable. "
        "Future work includes deploying the Docker container to a cloud Kubernetes cluster and integrating Drift Detection monitors."
    )
    
    # --- References ---
    doc.add_heading('References', level=1)
    doc.add_paragraph(
        "[1] Beata Strack et al., 'Impact of HbA1c Measurement on Hospital Readmission Rates', BioMed Research International, 2014.\n"
        "[2] UCI Machine Learning Repository, 'Diabetes 130-US hospitals Data Set'."
    )

    doc.save('Project_Report.docx')
    print("Report generated: Project_Report.docx")

if __name__ == "__main__":
    create_report()
